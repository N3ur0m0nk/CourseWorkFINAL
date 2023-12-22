import docx


def replace_keywords_in_docx(docx_path, keyword_dict, output_path):
    doc = docx.Document(docx_path)
    slovar = keyword_dict

    for paragraph in doc.paragraphs:
        for k, v in slovar.items():
            if k in paragraph.text:
                paragraph.text = paragraph.text.replace(k, v)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for k, v in slovar.items():
                    if k in cell.text:
                        cell.text = cell.text.replace(k, v)
    doc.save(output_path)


def FIOSplit(name):
    parts = name.split()
    return f'{parts[0]} {parts[1][0]}.{parts[2][0]}.'


keyword_dict = {
    "${Date}": "Значение1",
    "${FullFIO}": "Значение2",
    "${Specialty}": "Значение3",
    "${FullKafedra}": "Значение4",
    "${Kafedra}": "Значение5",
    "${Birthday}": "Значение6",
    "${Passport_Number}": "Значение7",
    "${Passport_Vidacha}": "Значение8",
    "${Passport_Date}": "Значение9",
    "${Registration_Address}": "Значение10",
    "${FIO}": "Значение11"
}
input_file = "Ученический договор Шаблон.docx"
output_file = "Ученический договор ИТОГ.docx"
